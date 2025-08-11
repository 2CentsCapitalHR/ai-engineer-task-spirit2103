import io
import json
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from chains.doc_type_chain import build_doc_type_chain
from chains.review_chain import build_review_chain
from chains.compliance_chain import build_compliance_chain
from utils import compare_against_checklist, CHECKLISTS
from langchain_huggingface import HuggingFaceEmbeddings

# === Build chains once ===
doc_type_chain = build_doc_type_chain()
review_chain = build_review_chain()
compliance_chain = build_compliance_chain()
embeddings_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")


def _safe_parse_json(maybe_json_str):
    """Try to parse a JSON string, else return None."""
    try:
        return json.loads(maybe_json_str)
    except Exception:
        import re
        m = re.search(r'(\[.*\]|\{.*\})', maybe_json_str, re.S)
        if m:
            try:
                return json.loads(m.group(1))
            except Exception:
                return None
        return None


def _add_comment_below_paragraph(paragraph, comment_text):
    """Add AI comment as a new styled paragraph below matched section."""
    comment_para = paragraph.insert_paragraph_before(f"[AI COMMENT] {comment_text}")
    run = comment_para.runs[0]
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 102, 204)  # blue


def _find_best_matching_paragraph(doc, target_text):
    """Find paragraph most semantically similar to target_text using embeddings."""
    if not target_text or len(target_text.strip()) < 5:
        return None
    para_texts = [p.text for p in doc.paragraphs if p.text.strip()]
    if not para_texts:
        return None
    target_emb = embeddings_model.embed_query(target_text)
    para_embs = embeddings_model.embed_documents(para_texts)
    sims = np.dot(para_embs, target_emb) / (
        np.linalg.norm(para_embs, axis=1) * np.linalg.norm(target_emb)
    )
    if len(sims) == 0:
        return None
    best_idx = int(np.argmax(sims))
    if sims[best_idx] < 0.4:  # similarity threshold
        return None
    return doc.paragraphs[best_idx]


def analyze_and_comment_docx_batch(docx_bytes_list, uploaded_filenames=None):
    if uploaded_filenames is None:
        uploaded_filenames = [f"doc_{i+1}.docx" for i in range(len(docx_bytes_list))]

    per_file_info = []
    detected_types = []

    for idx, doc_bytes in enumerate(docx_bytes_list):
        filename = uploaded_filenames[idx]
        doc = Document(io.BytesIO(doc_bytes))
        full_text = "\n".join([p.text for p in doc.paragraphs])

        # === 1. Document type detection (Embeddings only, no keyword fallback) ===
        try:
            ai_doc_type = doc_type_chain.invoke(full_text[:2000])
            if isinstance(ai_doc_type, dict):
                ai_doc_type = ai_doc_type.get("result") or ai_doc_type.get("answer") or str(ai_doc_type)
            ai_doc_type = str(ai_doc_type).strip().strip('"')
        except Exception:
            ai_doc_type = None

        if not ai_doc_type:
            ai_doc_type = "Classification Error"

        detected_types.append(ai_doc_type)

        # === 2. Review chain (structured JSON) ===
        try:
            review_prompt = f"Review the following {ai_doc_type} for completeness and correctness:\n{full_text[:4000]}"
            review_result = review_chain.invoke(review_prompt)
            if isinstance(review_result, dict):
                review_result = review_result.get("result") or review_result.get("answer") or str(review_result)
            review_json = _safe_parse_json(str(review_result))
            if not review_json:
                review_json = {"summary": str(review_result), "recommendations": []}
        except Exception as e:
            review_json = {"summary": f"Error generating review: {e}", "recommendations": []}

        # === 3. Compliance chain (strict JSON) ===
        try:
            compliance_prompt = f"Check compliance for {ai_doc_type}:\n{full_text[:4000]}"
            compliance_result = compliance_chain.invoke(compliance_prompt)
            if isinstance(compliance_result, dict):
                compliance_result = compliance_result.get("result") or compliance_result.get("answer") or str(compliance_result)
            compliance_json = _safe_parse_json(str(compliance_result))
            if not compliance_json:
                compliance_json = [{
                    "section": "N/A",
                    "issue": str(compliance_result),
                    "severity": "Low",
                    "suggestion": ""
                }]
        except Exception as e:
            compliance_json = [{
                "section": "N/A",
                "issue": f"Compliance check failed: {e}",
                "severity": "Low",
                "suggestion": ""
            }]

        per_file_info.append({
            "filename": filename,
            "chosen_type": ai_doc_type,
            "review_json": review_json,
            "compliance_issues": compliance_json,
            "doc_object": doc
        })

    # === 4. Determine process type ===
    process_scores = {}
    for proc, required_list in CHECKLISTS.items():
        score = sum(1 for r in required_list if r in detected_types)
        process_scores[proc] = score

    chosen_process = max(process_scores.items(), key=lambda x: (x[1], -len(CHECKLISTS.get(x[0], []))))[0]
    required_docs = CHECKLISTS.get(chosen_process, [])
    uploaded_docs_count = len(docx_bytes_list)
    checklist_result = compare_against_checklist(chosen_process, detected_types)
    missing_docs = checklist_result.get("missing_documents", [])

    # === 5. Aggregate issues & reviews ===
    issues_found = []
    review_summaries = []
    for pf in per_file_info:
        for issue in pf["compliance_issues"]:
            issues_found.append({
                "document": pf["chosen_type"],
                "filename": pf["filename"],
                "section": issue.get("section", "N/A"),
                "issue": issue.get("issue", ""),
                "severity": issue.get("severity", "Low"),
                "suggestion": issue.get("suggestion", "")
            })
        review_summaries.append({
            "document": pf["chosen_type"],
            "filename": pf["filename"],
            "summary": pf["review_json"].get("summary", ""),
            "recommendations": pf["review_json"].get("recommendations", [])
        })

    # === 6. Annotate DOCX files ===
    reviewed_docs = []
    for pf in per_file_info:
        doc = pf["doc_object"]

        header_para = doc.add_paragraph("=== AI REVIEW SUMMARY ===")
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        header_para.runs[0].bold = True
        doc.add_paragraph(f"Document Type: {pf['chosen_type']}")
        doc.add_paragraph("Summary:")
        doc.add_paragraph(f"- {pf['review_json'].get('summary', '')}")

        doc.add_paragraph(f"Issues Found ({len(pf['compliance_issues'])}):")
        for idx, issue in enumerate(pf["compliance_issues"], start=1):
            doc.add_paragraph(f"{idx}. {issue.get('issue', '')}")

        if pf["review_json"].get("recommendations"):
            doc.add_paragraph("Recommendations:")
            for rec in pf["review_json"]["recommendations"]:
                doc.add_paragraph(f"- {rec}")

        for issue in pf["compliance_issues"]:
            sec = issue.get("section", "")
            issue_text = issue.get("issue", "")
            suggestion = issue.get("suggestion", "")
            matched_para = None
            if sec and sec != "N/A":
                matched_para = _find_best_matching_paragraph(doc, sec)
            if not matched_para and issue_text:
                matched_para = _find_best_matching_paragraph(doc, issue_text)
            comment_text = f"{issue_text} | Suggestion: {suggestion}"
            if matched_para:
                _add_comment_below_paragraph(matched_para, comment_text)
            else:
                _add_comment_below_paragraph(doc.paragraphs[0], comment_text)

        out_stream = io.BytesIO()
        doc.save(out_stream)
        out_stream.seek(0)
        reviewed_docs.append((pf["filename"], out_stream.read()))

    combined_report = {
        "process": chosen_process,
        "documents_uploaded": uploaded_docs_count,
        "required_documents": len(required_docs),
        "missing_documents": missing_docs,
        "issues_found": issues_found,
        "reviews": review_summaries
    }

    return reviewed_docs, combined_report
